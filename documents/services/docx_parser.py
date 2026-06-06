import logging
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract full text from a docx file.
    Reads both paragraphs and tables.
    """
    try:
        doc = DocxDocument(file_path)
    except Exception as e:
        logger.error(f"Error opening docx file: {e}")
        raise ValueError(f"Invalid docx file: {e}")

    parts = []

    # 1) Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # 2) Tables (many technical documents contain tables)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    full_text = "\n".join(parts)

    if not full_text.strip():
        logger.warning("Docx file is empty or no text could be extracted.")

    return full_text